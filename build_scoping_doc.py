#!/usr/bin/env python3
"""Build the TC website next-phase scoping document. Plain, black, simple."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BLACK = RGBColor(0, 0, 0)
FONT = "Open Sans"

doc = Document()
n = doc.styles["Normal"]; n.font.name = FONT; n.font.size = Pt(11); n.font.color.rgb = BLACK
rpr = n.element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
rf.set(qn("w:ascii"), FONT); rf.set(qn("w:hAnsi"), FONT)

def title(t):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
    r = p.add_run(t); r.font.name = FONT; r.font.size = Pt(18); r.font.bold = True; r.font.color.rgb = BLACK
def heading(t):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t); r.font.name = FONT; r.font.size = Pt(13.5); r.font.bold = True; r.font.color.rgb = BLACK
def body(t, sa=8, bold=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(sa); p.paragraph_format.line_spacing = 1.15
    r = p.add_run(t); r.font.name = FONT; r.font.size = Pt(11); r.font.bold = bold; r.font.color.rgb = BLACK
def bullet(t, lead=None):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
    if lead:
        r = p.add_run(lead); r.font.name = FONT; r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = BLACK
    r = p.add_run(t); r.font.name = FONT; r.font.size = Pt(11); r.font.color.rgb = BLACK
def cell(c, text, bold=False, center=False):
    c.text = ""; p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.name = FONT; r.font.size = Pt(10); r.font.bold = bold; r.font.color.rgb = BLACK
def table(headers, rows, widths):
    tb = doc.add_table(rows=0, cols=len(headers)); tb.style = "Table Grid"
    hc = tb.add_row().cells
    for c, t in zip(hc, headers): cell(c, t, bold=True)
    for row in rows:
        cc = tb.add_row().cells
        for i, t in enumerate(row): cell(cc[i], t, center=(i == len(row) - 1 and len(row) > 1))
    for r in tb.rows:
        for i, c in enumerate(r.cells): c.width = Inches(widths[i])
    return tb

# ---- Title ----
title("Tracking California Website: Scope and Next Steps")
body("Live demo: https://jo-wilkin-tc.github.io/tc-website-redesign/", sa=12)

# ---- What we have done ----
heading("What we have done")
body("A demo of the redesigned website is live. It is built in plain HTML and CSS and shows the full design and structure. It contains:")
bullet("Home page: hero, two pillars (data and community), latest work, partner map, an interactive data/tool finder, publications, and topics.")
bullet("Tools page: embedded Data Explorer, current tools, and planned tools.")
bullet("Our Data page: data insights (sample dashboard), data briefs with filters, and open-science repositories.")
bullet("Our Communities page: interactive California map of community partners.")
bullet("Topics page (all 22 topics) and a topic-page template.")
bullet("Resources page populated with real publications, newsletters, videos, and news.")
bullet("15 individual news article pages generated from the existing site content.")
bullet("About page with the real staff list and an interactive partner directory.")
bullet("Working site search (Pagefind).")
body("The demo is not yet content-managed, fully migrated, bilingual, or on production hosting. Those are the tasks below.")

# ---- What needs to be done ----
heading("What needs to be done")
body("The following is the work to take the demo to a live, editable, maintained site. Hours are developer-hours and shown as ranges.")
table(
    ["Task", "What it involves", "Hours"],
    [
        ["1. Decisions and setup", "Confirm hosting, framework, editors, and IA. Short framework test (Next.js vs Astro). Set up the repository.", "16-28"],
        ["2. Rebuild in a framework", "Rebuild the demo in Next.js or Astro so it can support a content management system. Port all page templates and components.", "70-110"],
        ["3. Add TinaCMS", "Set up TinaCMS so staff can edit pages directly in the browser. Define the content structure for each content type. Add media handling and previews.", "45-75"],
        ["4. Migrate the content", "Move the existing site content (about 285 items) into the new structure: pages (nested up to 12 levels), news, publications, newsletters, videos, and projects. Check and correct.", "45-85"],
        ["5. Add strategic-planning content", "Update the site structure, navigation, and copy to reflect the team's strategic-planning work.", "20-36"],
        ["6. Search", "Move the demo search into the build so it updates automatically. Style the results.", "6-12"],
        ["7. Accessibility, SEO, testing", "Meet accessibility standards (WCAG / Section 508). Keep existing page addresses working (redirects). Add sitemap and link previews. Test across browsers and devices.", "30-55"],
        ["8. Update process and training", "Write the process for updating the site. Write a short maintainer guide. Train staff.", "16-28"],
        ["9. Hosting and launch", "Set up hosting and automated deployment. Point the domain at the new site. Launch.", "16-30"],
        ["Project management and contingency", "Coordination, reviews, and unforeseen work (about 15%).", "40-70"],
    ],
    [1.9, 4.0, 0.9],
)
body("")
body("Total: about 300 to 530 hours (roughly 8 to 14 weeks of focused work for one person).", bold=True)
body("Optional, decide separately: Spanish / bilingual support adds about 30 to 55 hours. The current site has Spanish content; it is cheaper to plan for now than to add later.")

# ---- Timeline ----
heading("Timeline")
body("Order of work. Calendar time assumes one developer plus part-time review; it shortens with more people.")
table(
    ["Phase", "Focus", "Weeks"],
    [
        ["0", "Decisions, framework test, review strategy work", "1-2"],
        ["1", "Rebuild in framework, templates and components", "3-4"],
        ["2", "TinaCMS, content structure, previews", "3-4"],
        ["3", "Content migration and strategic-planning content", "3-5"],
        ["4", "Search, accessibility, SEO, testing (and bilingual, if included)", "2-4"],
        ["5", "Hosting, launch, staff training", "1-2"],
    ],
    [0.7, 5.2, 0.9],
)
body("")
body("Total calendar time: about 13 to 21 weeks part-time, or less with dedicated time.", bold=True)

# ---- Decisions needed ----
heading("Decisions needed before we start")
bullet("Bitbucket or GitHub. Still open; some health-department laptops cannot use GitHub. This decides the repository and deployment.", lead="Where the code lives. ")
bullet("Next.js or Astro. A short test in Phase 0 will confirm.", lead="Framework. ")
bullet("In the first version, or added later.", lead="Bilingual. ")
bullet("Who will edit the site (names). This shapes the CMS setup and training.", lead="Editors. ")
bullet("Share the strategic-planning documents so this work can be scoped accurately.", lead="Strategy work. ")
bullet("Confirm the data team is willing to share dataset and repository links before they go on the site.", lead="Data team. ")

# ---- Update process ----
heading("How the site will be updated")
bullet("Staff edit pages in TinaCMS in the browser (text, a news post, a new publication). Changes are saved, previewed, then published. No code required.", lead="Everyday content. ")
bullet("A developer handles new sections, structure changes, and bulk updates through the repository.", lead="Structural changes. ")
bullet("Newsletters continue through Constant Contact. Data updates follow the data team's schedule. Content reviewed quarterly.", lead="Ongoing. ")
bullet("Name one content owner and one technical maintainer, with a written guide, so the site does not depend on a single person.", lead="Roles. ")

import os
out = "/Users/joannawilkin/TC-projects/tc-website-redesign/docs/TC-website-next-phase-scoping.docx"
os.makedirs(os.path.dirname(out), exist_ok=True); doc.save(out); print("Saved:", out)
